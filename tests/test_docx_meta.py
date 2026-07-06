import json
import os
import hashlib
import tempfile
import zipfile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree
from PIL import Image as PILImage

from md2word.meta import (
    DocxMetaExtractor,
    DocxMetaRenderer,
    verify_docx_metadata_roundtrip,
)


def _part_bytes(docx_path, name):
    with zipfile.ZipFile(docx_path, "r") as package:
        return package.read(name)


def _canonical_xml(data):
    root = etree.fromstring(data)
    return etree.tostring(root, method="c14n")


def _rewrite_docx(source_path, target_path, replacements):
    with zipfile.ZipFile(source_path, "r") as src, zipfile.ZipFile(
        target_path, "w", compression=zipfile.ZIP_DEFLATED
    ) as dst:
        for info in src.infolist():
            data = replacements.get(info.filename, src.read(info.filename))
            new_info = zipfile.ZipInfo(info.filename, date_time=info.date_time)
            new_info.compress_type = info.compress_type
            new_info.external_attr = info.external_attr
            new_info.internal_attr = info.internal_attr
            new_info.create_system = info.create_system
            new_info.create_version = info.create_version
            new_info.extract_version = info.extract_version
            new_info.flag_bits = info.flag_bits
            new_info.volume = info.volume
            new_info.extra = info.extra
            new_info.comment = info.comment
            dst.writestr(new_info, data)
        dst.comment = src.comment


def _replace_first_xml_text(node, old, new):
    if node.get("kind") == "element":
        tag = node.get("tag", {})
        if tag.get("display") == "w:t" and node.get("text") == old:
            node["text"] = new
            return True
    for child in node.get("children", []):
        if _replace_first_xml_text(child, old, new):
            return True
    return False


def test_docx_metadata_extracts_xml_tree_and_resources():
    with tempfile.TemporaryDirectory() as tmp:
        img_path = os.path.join(tmp, "image.png")
        PILImage.new("RGB", (10, 10), color="blue").save(img_path)

        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("Colored text")
        run.font.color.rgb = RGBColor(0x12, 0x34, 0x56)
        doc.add_picture(img_path)
        docx_path = os.path.join(tmp, "sample.docx")
        doc.save(docx_path)

        meta_dir = os.path.join(tmp, "meta")
        metadata = DocxMetaExtractor().extract(docx_path, meta_dir)

        assert metadata["format"] == "md2word.docx-package-meta"
        assert os.path.exists(os.path.join(meta_dir, "document.json"))

        document_entry = next(e for e in metadata["entries"] if e["path"] == "word/document.xml")
        assert document_entry["kind"] == "xml"
        assert document_entry["xml"]["root"]["tag"]["display"] == "w:document"

        media_entries = [e for e in metadata["entries"] if e["path"].startswith("word/media/")]
        assert len(media_entries) == 1
        assert media_entries[0]["kind"] == "resource"
        assert os.path.exists(os.path.join(meta_dir, media_entries[0]["resource"]))

        with open(os.path.join(meta_dir, "document.json"), "r", encoding="utf-8") as f:
            persisted = json.load(f)
        assert persisted["entries"][0]["path"] == metadata["entries"][0]["path"]
        assert "document" in persisted


def test_docx_metadata_restores_openable_package_with_equal_xml_parts():
    with tempfile.TemporaryDirectory() as tmp:
        doc = Document()
        doc.add_heading("Roundtrip", level=1)
        p = doc.add_paragraph("Before ")
        bold = p.add_run("bold")
        bold.bold = True
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"
        table.cell(1, 0).text = "C"
        table.cell(1, 1).text = "D"

        source_docx = os.path.join(tmp, "source.docx")
        restored_docx = os.path.join(tmp, "restored.docx")
        meta_dir = os.path.join(tmp, "meta")
        doc.save(source_docx)

        DocxMetaExtractor().extract(source_docx, meta_dir)
        DocxMetaRenderer().restore(meta_dir, restored_docx)

        source_bytes = open(source_docx, "rb").read()
        restored_bytes = open(restored_docx, "rb").read()
        assert hashlib.sha256(restored_bytes).hexdigest() == hashlib.sha256(source_bytes).hexdigest()
        assert restored_bytes == source_bytes

        restored = Document(restored_docx)
        assert restored.paragraphs[0].text == "Roundtrip"
        assert restored.paragraphs[1].text == "Before bold"
        assert restored.tables[0].cell(1, 1).text == "D"

        assert _canonical_xml(_part_bytes(source_docx, "word/document.xml")) == (
            _canonical_xml(_part_bytes(restored_docx, "word/document.xml"))
        )
        assert _canonical_xml(_part_bytes(source_docx, "word/styles.xml")) == (
            _canonical_xml(_part_bytes(restored_docx, "word/styles.xml"))
        )


def test_docx_metadata_preserves_xml_declaration_preamble():
    with tempfile.TemporaryDirectory() as tmp:
        doc = Document()
        doc.add_paragraph("Declaration preamble")
        source_docx = os.path.join(tmp, "source.docx")
        mutated_docx = os.path.join(tmp, "mutated.docx")
        restored_docx = os.path.join(tmp, "restored.docx")
        doc.save(source_docx)

        content_types = _part_bytes(source_docx, "[Content_Types].xml")
        if b"?>" in content_types:
            content_types = content_types.split(b"?>", 1)[1].lstrip(b"\r\n")
        content_types = (
            b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\r\n'
            + content_types
        )
        _rewrite_docx(source_docx, mutated_docx, {"[Content_Types].xml": content_types})

        metadata = DocxMetaExtractor().extract(mutated_docx, os.path.join(tmp, "meta"))
        content_types_entry = next(
            entry for entry in metadata["entries"] if entry["path"] == "[Content_Types].xml"
        )
        assert content_types_entry["xml"]["preamble"].endswith("?>\r\n")

        DocxMetaRenderer().restore(os.path.join(tmp, "meta"), restored_docx)

        assert _part_bytes(mutated_docx, "[Content_Types].xml") == _part_bytes(
            restored_docx, "[Content_Types].xml"
        )


def test_docx_metadata_edits_fall_back_from_exact_payload_cache():
    with tempfile.TemporaryDirectory() as tmp:
        doc = Document()
        doc.add_paragraph("Original")
        source_docx = os.path.join(tmp, "source.docx")
        restored_docx = os.path.join(tmp, "restored.docx")
        meta_dir = os.path.join(tmp, "meta")
        doc.save(source_docx)

        metadata = DocxMetaExtractor().extract(source_docx, meta_dir)
        document_entry = next(
            entry for entry in metadata["entries"] if entry["path"] == "word/document.xml"
        )
        assert _replace_first_xml_text(document_entry["xml"]["root"], "Original", "Changed")
        with open(os.path.join(meta_dir, "document.json"), "w", encoding="utf-8") as f:
            json.dump(metadata, f, ensure_ascii=False, indent=2)
            f.write("\n")

        DocxMetaRenderer().restore(meta_dir, restored_docx)

        restored = Document(restored_docx)
        assert restored.paragraphs[0].text == "Changed"
        assert open(restored_docx, "rb").read() != open(source_docx, "rb").read()


def test_docx_metadata_semantic_index_covers_priority_formatting():
    with tempfile.TemporaryDirectory() as tmp:
        img_path = os.path.join(tmp, "image.png")
        PILImage.new("RGB", (16, 12), color="green").save(img_path)

        doc = Document()
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.space_after = Pt(18)
        run = paragraph.add_run("Styled")
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x12, 0x34, 0x56)
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

        doc.add_picture(img_path, width=Inches(1))

        table = doc.add_table(rows=2, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"
        table.cell(0, 0).merge(table.cell(0, 1))
        table.cell(1, 0).text = "C"
        table.cell(1, 1).text = "D"
        table.cell(1, 1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        docx_path = os.path.join(tmp, "semantic.docx")
        doc.save(docx_path)

        metadata = DocxMetaExtractor().extract(docx_path, os.path.join(tmp, "meta"))
        body = metadata["document"]["body"]

        first_paragraph = body[0]
        assert first_paragraph["properties"]["alignment"] == "center"
        assert first_paragraph["properties"]["indent"]["w:left"] == "360"
        assert first_paragraph["properties"]["spacing"]["w:after"] == "360"

        run_props = first_paragraph["runs"][0]["properties"]
        assert run_props["bold"] is True
        assert run_props["fonts"]["w:ascii"] == "Arial"
        assert run_props["size_half_points"] == "28"
        assert run_props["color"]["w:val"] == "123456"
        assert run_props["highlight"]["w:val"] == "yellow"

        image_run = body[1]["runs"][0]
        drawing = image_run["drawings"][0]
        assert drawing["type"] == "inline"
        assert drawing["relationship"]["package_path"].startswith("word/media/")
        assert drawing["resource"].startswith("resources/word/media/")
        assert drawing["extent"]["cx"] == "914400"

        semantic_table = body[2]
        assert semantic_table["properties"]["alignment"] == "center"
        assert semantic_table["rows"][0]["cells"][0]["properties"]["grid_span"] == "2"
        assert semantic_table["rows"][1]["cells"][1]["properties"]["vertical_align"] == "center"

        section = body[-1]
        assert section["type"] == "section_properties"
        assert section["properties"]["page_margins"]["w:left"] == "1800"


def test_docx_metadata_roundtrip_verifier_reports_byte_identical():
    with tempfile.TemporaryDirectory() as tmp:
        doc = Document()
        doc.add_paragraph("Verification")
        source_docx = os.path.join(tmp, "verify.docx")
        doc.save(source_docx)

        result = verify_docx_metadata_roundtrip(source_docx, os.path.join(tmp, "roundtrip"))

        assert result["byte_identical"] is True
        assert result["source_sha256"] == result["restored_sha256"]
        assert os.path.exists(result["metadata_dir"])
        assert os.path.exists(result["restored"])


def test_docx_metadata_semantic_index_covers_package_support_parts():
    with tempfile.TemporaryDirectory() as tmp:
        doc = Document()
        doc.core_properties.title = "Semantic Metadata"
        doc.core_properties.author = "MD2Word"
        doc.add_heading("Heading", level=1)
        doc.add_paragraph("Numbered item", style="List Number")

        docx_path = os.path.join(tmp, "support-parts.docx")
        doc.save(docx_path)

        metadata = DocxMetaExtractor().extract(docx_path, os.path.join(tmp, "meta"))
        semantic = metadata["document"]

        assert semantic["metadata"]["core"]["title"]["text"] == "Semantic Metadata"
        assert semantic["metadata"]["core"]["creator"]["text"] == "MD2Word"
        assert semantic["metadata"]["app"]["Template"]["text"] == "Normal.dotm"

        assert semantic["settings"]["default_tab_stop"] == "720"
        assert semantic["settings"]["zoom"]["w:val"] == "bestFit"

        styles_by_id = {
            style["style_id"]: style
            for style in semantic["styles"]["styles"]
            if "style_id" in style
        }
        assert styles_by_id["Normal"]["type"] == "paragraph"
        assert styles_by_id["Heading1"]["name"] == "heading 1"

        numbering = semantic["numbering"]
        assert numbering["abstract_numbers"][0]["levels"][0]["format"] == "decimal"
        assert numbering["abstract_numbers"][0]["levels"][0]["text"] == "%1."
        assert numbering["numbers"][0]["abstract_num_id"] is not None

        font_names = {font["name"] for font in semantic["fonts"]["fonts"]}
        assert "Symbol" in font_names
        assert "Times New Roman" in font_names

        assert semantic["theme"]["name"] == "Office Theme"
        assert semantic["theme"]["colors"]["accent1"]["attributes"]["val"] == "4F81BD"
        assert semantic["theme"]["fonts"]["major"]["latin"]["typeface"] == "Calibri"


def test_docx_metadata_semantic_index_covers_fields_sdt_and_revisions():
    with tempfile.TemporaryDirectory() as tmp:
        doc = Document()

        field_p = doc.add_paragraph()
        begin = field_p.add_run()._r
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        begin.append(fld_begin)
        instr_run = field_p.add_run()._r
        instr = OxmlElement("w:instrText")
        instr.text = "PAGE"
        instr_run.append(instr)
        separate = field_p.add_run()._r
        fld_separate = OxmlElement("w:fldChar")
        fld_separate.set(qn("w:fldCharType"), "separate")
        separate.append(fld_separate)
        field_p.add_run("1")
        end = field_p.add_run()._r
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        end.append(fld_end)

        sdt = OxmlElement("w:sdt")
        sdt_pr = OxmlElement("w:sdtPr")
        alias = OxmlElement("w:alias")
        alias.set(qn("w:val"), "Client Name")
        tag = OxmlElement("w:tag")
        tag.set(qn("w:val"), "client_name")
        text = OxmlElement("w:text")
        sdt_pr.extend([alias, tag, text])
        sdt_content = OxmlElement("w:sdtContent")
        sdt_p = OxmlElement("w:p")
        sdt_r = OxmlElement("w:r")
        sdt_t = OxmlElement("w:t")
        sdt_t.text = "Ada"
        sdt_r.append(sdt_t)
        sdt_p.append(sdt_r)
        sdt_content.append(sdt_p)
        sdt.extend([sdt_pr, sdt_content])
        doc.element.body.insert(-1, sdt)

        rev_p = doc.add_paragraph()
        ins = OxmlElement("w:ins")
        ins.set(qn("w:id"), "1")
        ins.set(qn("w:author"), "Reviewer")
        ins.set(qn("w:date"), "2026-07-06T00:00:00Z")
        ins_run = OxmlElement("w:r")
        ins_text = OxmlElement("w:t")
        ins_text.text = "inserted"
        ins_run.append(ins_text)
        ins.append(ins_run)
        rev_p._p.append(ins)

        delete = OxmlElement("w:del")
        delete.set(qn("w:id"), "2")
        delete.set(qn("w:author"), "Reviewer")
        del_run = OxmlElement("w:r")
        del_text = OxmlElement("w:delText")
        del_text.text = "deleted"
        del_run.append(del_text)
        delete.append(del_run)
        rev_p._p.append(delete)

        docx_path = os.path.join(tmp, "advanced.docx")
        doc.save(docx_path)

        metadata = DocxMetaExtractor().extract(docx_path, os.path.join(tmp, "meta"))
        body = metadata["document"]["body"]

        field_para = body[0]
        assert field_para["fields"][0]["instruction"] == "PAGE"
        assert field_para["fields"][0]["result"] == "1"
        assert field_para["runs"][0]["field_chars"][0]["w:fldCharType"] == "begin"
        assert field_para["runs"][1]["field_instruction"] == "PAGE"

        sdt_block = next(element for element in body if element["type"] == "structured_document_tag")
        assert sdt_block["properties"]["alias"] == "Client Name"
        assert sdt_block["properties"]["tag"] == "client_name"
        assert sdt_block["properties"]["control"]["type"] == "text"
        assert sdt_block["text"] == "Ada"

        rev_para = body[-2]
        revisions = [run for run in rev_para["runs"] if run["type"] == "revision"]
        assert revisions[0]["kind"] == "ins"
        assert revisions[0]["author"] == "Reviewer"
        assert revisions[0]["text"] == "inserted"
        assert revisions[1]["kind"] == "del"
        assert revisions[1]["runs"][0]["deleted_text"] == "deleted"


def test_docx_metadata_semantic_index_covers_advanced_package_parts():
    with tempfile.TemporaryDirectory() as tmp:
        doc = Document()
        doc.add_paragraph("Advanced package parts")
        docx_path = os.path.join(tmp, "advanced-parts.docx")
        doc.save(docx_path)

        with zipfile.ZipFile(docx_path, "a", compression=zipfile.ZIP_DEFLATED) as package:
            package.writestr(
                "word/charts/chart1.xml",
                """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<c:chartSpace xmlns:c="http://schemas.openxmlformats.org/drawingml/2006/chart" xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
  <c:chart>
    <c:title><c:tx><c:rich><a:p><a:r><a:t>Sales</a:t></a:r></a:p></c:rich></c:tx></c:title>
    <c:plotArea>
      <c:barChart>
        <c:barDir val="col"/>
        <c:grouping val="clustered"/>
        <c:varyColors val="0"/>
        <c:ser>
          <c:idx val="0"/>
          <c:order val="0"/>
          <c:tx><c:strRef><c:f>Sheet1!$B$1</c:f><c:strCache><c:pt idx="0"><c:v>Revenue</c:v></c:pt></c:strCache></c:strRef></c:tx>
          <c:cat><c:strRef><c:f>Sheet1!$A$2:$A$3</c:f><c:strCache><c:pt idx="0"><c:v>Q1</c:v></c:pt><c:pt idx="1"><c:v>Q2</c:v></c:pt></c:strCache></c:strRef></c:cat>
          <c:val><c:numRef><c:f>Sheet1!$B$2:$B$3</c:f><c:numCache><c:pt idx="0"><c:v>10</c:v></c:pt><c:pt idx="1"><c:v>20</c:v></c:pt></c:numCache></c:numRef></c:val>
        </c:ser>
        <c:axId val="123"/>
        <c:axId val="456"/>
      </c:barChart>
      <c:catAx><c:axId val="123"/><c:axPos val="b"/><c:tickLblPos val="nextTo"/><c:crossAx val="456"/></c:catAx>
      <c:valAx><c:axId val="456"/><c:axPos val="l"/><c:numFmt formatCode="General" sourceLinked="1"/><c:crossAx val="123"/><c:crosses val="autoZero"/></c:valAx>
    </c:plotArea>
    <c:legend><c:legendPos val="r"/><c:overlay val="0"/></c:legend>
  </c:chart>
</c:chartSpace>
""",
            )
            package.writestr(
                "word/diagrams/data1.xml",
                """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<dgm:dataModel xmlns:dgm="http://schemas.openxmlformats.org/drawingml/2006/diagram" xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
  <dgm:ptLst>
    <dgm:pt modelId="0" type="doc">
      <dgm:prSet phldr="0"/>
    </dgm:pt>
    <dgm:pt modelId="1" type="node">
      <dgm:prSet custAng="0"/>
      <dgm:t><a:txBody><a:p><a:r><a:t>First node</a:t></a:r></a:p></a:txBody></dgm:t>
    </dgm:pt>
  </dgm:ptLst>
  <dgm:cxnLst>
    <dgm:cxn modelId="2" type="parOf" srcId="0" destId="1" srcOrd="0" destOrd="0"/>
  </dgm:cxnLst>
</dgm:dataModel>
""",
            )
            package.writestr(
                "word/diagrams/layout1.xml",
                """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<dgm:layoutDef xmlns:dgm="http://schemas.openxmlformats.org/drawingml/2006/diagram" uniqueId="layout-1" minVer="12.0">
  <dgm:title val="Basic process"/>
  <dgm:desc val="Process layout"/>
  <dgm:catLst><dgm:cat type="process" pri="100"/></dgm:catLst>
  <dgm:layoutNode name="node" styleLbl="node0">
    <dgm:alg type="lin"/>
    <dgm:shape type="rect"/>
    <dgm:presOf axis="self" ptType="node"/>
    <dgm:constrLst><dgm:constr type="w" val="100"/></dgm:constrLst>
    <dgm:ruleLst><dgm:rule type="primFontSz" val="1200"/></dgm:ruleLst>
    <dgm:varLst><dgm:dir val="norm"/></dgm:varLst>
  </dgm:layoutNode>
</dgm:layoutDef>
""",
            )
            package.writestr(
                "word/diagrams/quickStyle1.xml",
                """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<dgm:styleDef xmlns:dgm="http://schemas.openxmlformats.org/drawingml/2006/diagram" uniqueId="style-1" minVer="12.0">
  <dgm:title val="Simple style"/>
  <dgm:desc val="Simple"/>
  <dgm:catLst><dgm:cat type="simple" pri="100"/></dgm:catLst>
  <dgm:styleLbl name="node0"><dgm:style/></dgm:styleLbl>
</dgm:styleDef>
""",
            )
            package.writestr(
                "word/diagrams/colors1.xml",
                """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<dgm:colorsDef xmlns:dgm="http://schemas.openxmlformats.org/drawingml/2006/diagram" xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" uniqueId="colors-1" minVer="12.0">
  <dgm:title val="Colorful"/>
  <dgm:desc val="Colors"/>
  <dgm:catLst><dgm:cat type="colorful" pri="100"/></dgm:catLst>
  <dgm:styleLbl name="node0">
    <dgm:fillClrLst><a:srgbClr val="FF0000"/></dgm:fillClrLst>
    <dgm:linClrLst><a:srgbClr val="00FF00"/></dgm:linClrLst>
  </dgm:styleLbl>
</dgm:colorsDef>
""",
            )
            package.writestr(
                "word/activeX/activeX1.xml",
                """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<ax:ocx xmlns:ax="http://schemas.microsoft.com/office/2006/activeX" classid="{00000000-0000-0000-0000-000000000000}"/>
""",
            )
            package.writestr(
                "word/people.xml",
                """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:people xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:person w:name="Reviewer" w:id="person-1"/>
</w:people>
""",
            )
            package.writestr(
                "word/glossary/document.xml",
                """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:glossaryDocument xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:docParts>
    <w:docPart>
      <w:docPartPr><w:name w:val="Reusable block"/></w:docPartPr>
      <w:docPartBody><w:p><w:r><w:t>Glossary text</w:t></w:r></w:p></w:docPartBody>
    </w:docPart>
  </w:docParts>
</w:glossaryDocument>
""",
            )
            package.writestr(
                "customXml/item99.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<customer><name>Ada</name></customer>
""",
            )
            package.writestr("word/embeddings/oleObject1.bin", b"ole-binary")
            package.writestr("word/vbaProject.bin", b"vba-binary")
            package.writestr("word/activeX/activeX1.bin", b"activex-binary")

        metadata = DocxMetaExtractor().extract(docx_path, os.path.join(tmp, "meta"))
        semantic = metadata["document"]

        assert semantic["charts"][0]["package_path"] == "word/charts/chart1.xml"
        assert semantic["charts"][0]["chart_types"] == ["barChart"]
        assert semantic["charts"][0]["series_count"] == 1
        assert "Sales" in semantic["charts"][0]["title"]
        chart_group = semantic["charts"][0]["chart_groups"][0]
        assert chart_group["type"] == "barChart"
        assert chart_group["bar_direction"] == "col"
        assert chart_group["grouping"] == "clustered"
        assert chart_group["axis_ids"] == ["123", "456"]
        series = chart_group["series"][0]
        assert series["title"]["points"][0]["value"] == "Revenue"
        assert series["categories"]["formula"] == "Sheet1!$A$2:$A$3"
        assert series["categories"]["points"][1]["value"] == "Q2"
        assert series["values"]["formula"] == "Sheet1!$B$2:$B$3"
        assert series["values"]["points"][0]["value"] == "10"
        assert semantic["charts"][0]["axes"][0]["type"] == "catAx"
        assert semantic["charts"][0]["axes"][1]["number_format"]["formatCode"] == "General"
        assert semantic["charts"][0]["legend"]["position"] == "r"

        diagram = semantic["diagrams"]["data"][0]
        assert diagram["root"] == "dataModel"
        assert diagram["points"][1]["text"] == "First node"
        assert diagram["connections"][0]["source_id"] == "0"
        assert diagram["connections"][0]["destination_id"] == "1"
        layout = semantic["diagrams"]["layouts"][0]
        assert layout["title"] == "Basic process"
        assert layout["layout_nodes"][0]["algorithm"]["attributes"]["type"] == "lin"
        assert layout["layout_nodes"][0]["shape"]["attributes"]["type"] == "rect"
        style = semantic["diagrams"]["quick_styles"][0]
        assert style["style_labels"][0]["name"] == "node0"
        colors = semantic["diagrams"]["colors"][0]
        assert colors["color_labels"][0]["fill_colors"][0]["attributes"]["val"] == "FF0000"
        assert colors["color_labels"][0]["line_colors"][0]["attributes"]["val"] == "00FF00"
        assert semantic["embeddings"][0]["package_path"] == "word/embeddings/oleObject1.bin"
        assert semantic["vba_project"]["package_path"] == "word/vbaProject.bin"
        assert semantic["active_x"]["controls"][0]["package_path"] == "word/activeX/activeX1.xml"
        assert semantic["active_x"]["binaries"][0]["package_path"] == "word/activeX/activeX1.bin"
        assert semantic["people"][0]["attributes"]["w:name"] == "Reviewer"
        assert semantic["glossary"]["doc_parts"][0]["name"] == "Reusable block"
        assert semantic["glossary"]["doc_parts"][0]["body"][0]["text"] == "Glossary text"
        custom_xml = next(item for item in semantic["custom_xml"] if item["package_path"] == "customXml/item99.xml")
        assert custom_xml["root"] == "customer"
        assert custom_xml["text"] == "Ada"


def test_docx_metadata_semantic_index_covers_vml_shapes_and_textboxes():
    with tempfile.TemporaryDirectory() as tmp:
        doc = Document()
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()

        pict = OxmlElement("w:pict")
        v_ns = "urn:schemas-microsoft-com:vml"
        shape = etree.Element(f"{{{v_ns}}}shape", nsmap={"v": v_ns})
        shape.set("id", "TextBox1")
        shape.set("type", "#_x0000_t202")
        shape.set("style", "position:absolute;margin-left:10pt;margin-top:20pt;width:120pt;height:40pt")
        shape.set("fillcolor", "#FFFF00")
        shape.set("strokecolor", "#FF0000")
        fill = etree.Element(f"{{{v_ns}}}fill")
        fill.set("color2", "#00FF00")
        fill.set("type", "gradient")
        stroke = etree.Element(f"{{{v_ns}}}stroke")
        stroke.set("weight", "1pt")
        stroke.set("dashstyle", "dash")
        textbox = etree.Element(f"{{{v_ns}}}textbox")
        textbox.set("inset", "2pt,2pt,2pt,2pt")
        txbx = OxmlElement("w:txbxContent")
        txbx_p = OxmlElement("w:p")
        txbx_r = OxmlElement("w:r")
        txbx_t = OxmlElement("w:t")
        txbx_t.text = "Inside shape"
        txbx_r.append(txbx_t)
        txbx_p.append(txbx_r)
        txbx.append(txbx_p)
        textbox.append(txbx)
        shape.extend([fill, stroke, textbox])
        pict.append(shape)
        run._r.append(pict)

        docx_path = os.path.join(tmp, "vml.docx")
        doc.save(docx_path)

        metadata = DocxMetaExtractor().extract(docx_path, os.path.join(tmp, "meta"))
        run_meta = metadata["document"]["body"][0]["runs"][0]
        vml = run_meta["vml"][0]

        assert vml["type"] == "vml_shape"
        assert vml["shape_type"] == "shape"
        assert vml["attributes"]["id"] == "TextBox1"
        assert vml["style"]["position"] == "absolute"
        assert vml["style"]["width"] == "120pt"
        assert vml["fill"]["attributes"]["type"] == "gradient"
        assert vml["stroke"]["attributes"]["dashstyle"] == "dash"
        assert vml["textbox"]["attributes"]["inset"] == "2pt,2pt,2pt,2pt"
        assert vml["textbox"]["text"] == "Inside shape"
        assert vml["textbox"]["body"][0]["text"] == "Inside shape"


def test_docx_metadata_semantic_index_covers_drawingml_picture_visual_properties():
    with tempfile.TemporaryDirectory() as tmp:
        img_path = os.path.join(tmp, "image.png")
        PILImage.new("RGB", (20, 10), color="purple").save(img_path)

        doc = Document()
        doc.add_picture(img_path, width=Inches(1))
        source_docx = os.path.join(tmp, "picture-source.docx")
        target_docx = os.path.join(tmp, "picture-mutated.docx")
        doc.save(source_docx)

        ns = {
            "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
            "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
            "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
            "rel": "http://schemas.openxmlformats.org/package/2006/relationships",
        }
        with zipfile.ZipFile(source_docx, "r") as package:
            document_xml = etree.fromstring(package.read("word/document.xml"))
            rels_xml = etree.fromstring(package.read("word/_rels/document.xml.rels"))

        blip_fill = document_xml.find(".//pic:blipFill", namespaces=ns)
        src_rect = etree.Element(f"{{{ns['a']}}}srcRect")
        src_rect.set("l", "1000")
        src_rect.set("t", "2000")
        src_rect.set("r", "3000")
        src_rect.set("b", "4000")
        blip_fill.append(src_rect)

        xfrm = document_xml.find(".//pic:spPr/a:xfrm", namespaces=ns)
        xfrm.set("rot", "600000")
        xfrm.set("flipH", "1")

        sp_pr = document_xml.find(".//pic:spPr", namespaces=ns)
        line = etree.Element(f"{{{ns['a']}}}ln")
        line.set("w", "12700")
        solid = etree.SubElement(line, f"{{{ns['a']}}}solidFill")
        color = etree.SubElement(solid, f"{{{ns['a']}}}srgbClr")
        color.set("val", "FF0000")
        dash = etree.SubElement(line, f"{{{ns['a']}}}prstDash")
        dash.set("val", "dash")
        sp_pr.append(line)

        effects = etree.Element(f"{{{ns['a']}}}effectLst")
        glow = etree.SubElement(effects, f"{{{ns['a']}}}glow")
        glow.set("rad", "40000")
        glow_color = etree.SubElement(glow, f"{{{ns['a']}}}srgbClr")
        glow_color.set("val", "00FF00")
        sp_pr.append(effects)

        c_nv_pr = document_xml.find(".//pic:cNvPr", namespaces=ns)
        hlink = etree.SubElement(c_nv_pr, f"{{{ns['a']}}}hlinkClick")
        hlink.set(f"{{{ns['r']}}}id", "rIdPictureLink")
        hlink.set("tooltip", "Open example")

        rel = etree.SubElement(rels_xml, f"{{{ns['rel']}}}Relationship")
        rel.set("Id", "rIdPictureLink")
        rel.set("Type", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink")
        rel.set("Target", "https://example.com")
        rel.set("TargetMode", "External")

        _rewrite_docx(
            source_docx,
            target_docx,
            {
                "word/document.xml": etree.tostring(
                    document_xml, xml_declaration=True, encoding="UTF-8", standalone=True
                ),
                "word/_rels/document.xml.rels": etree.tostring(
                    rels_xml, xml_declaration=True, encoding="UTF-8", standalone=True
                ),
            },
        )

        metadata = DocxMetaExtractor().extract(target_docx, os.path.join(tmp, "meta"))
        drawing = metadata["document"]["body"][0]["runs"][0]["drawings"][0]
        picture = drawing["picture"]

        assert picture["crop"] == {"l": "1000", "t": "2000", "r": "3000", "b": "4000"}
        assert picture["transform"]["rotation"] == "600000"
        assert picture["transform"]["flip_horizontal"] is True
        assert picture["line"]["attributes"]["w"] == "12700"
        assert picture["line"]["solid_fill"]["colors"][0]["attributes"]["val"] == "FF0000"
        assert picture["line"]["dash"]["val"] == "dash"
        assert picture["effects"]["effectLst"][0]["type"] == "glow"
        assert picture["effects"]["effectLst"][0]["attributes"]["rad"] == "40000"
        assert picture["hyperlink"]["relationship"]["target"] == "https://example.com"
        assert picture["hyperlink"]["tooltip"] == "Open example"
