import os
import tempfile

from docx import Document

from md2word.html import render_docx_to_html
from md2word.meta import DocxMetaExtractor


def test_full_coverage_fixture_exists_and_renders_to_html():
    md_path = os.path.join("examples", "full_coverage.md")
    docx_path = os.path.join("examples", "full_coverage.docx")
    assert os.path.exists(md_path)
    assert os.path.exists(docx_path)

    doc = Document(docx_path)
    assert len(doc.paragraphs) >= 50
    assert len(doc.tables) >= 2
    assert len(doc.sections) >= 2

    with tempfile.TemporaryDirectory() as tmp:
        metadata = DocxMetaExtractor().extract(docx_path, os.path.join(tmp, "meta"))
        semantic = metadata["document"]

        assert semantic["headers_footers"]["headers"]
        assert semantic["headers_footers"]["footers"]
        assert semantic["footnotes"][0]["text"] == "Issue 43 full coverage footnote."
        assert semantic["endnotes"][0]["text"] == "Issue 43 full coverage endnote."
        assert semantic["comments"][0]["text"] == "Issue 43 full coverage comment."
        assert len(semantic["sections"]) >= 2

        html_path = os.path.join(tmp, "full_coverage.html")
        render_docx_to_html(docx_path, html_path)
        html = open(html_path, "r", encoding="utf-8").read()

    assert "MathJax" in html
    assert "data:image/png;base64," in html
    assert "First page header" in html
    assert "Issue 43 full coverage comment" in html
    assert "Issue43Bookmark" in html
    assert 'colspan="2"' in html
