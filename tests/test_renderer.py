import os
import tempfile

from docx import Document as DocxDocument

from md2word.renderer.docx_renderer import DocxRenderer
from md2word.model.document import (
    Document, Heading, Paragraph, TextRun,
    CodeBlock, ListBlock, ListItem, Table, Image, HorizontalRule,
)


def _render_and_check(elements) -> DocxDocument:
    doc_model = Document(elements=elements)
    renderer = DocxRenderer()
    tmp = tempfile.mktemp(suffix=".docx")
    try:
        renderer.render(doc_model, tmp)
        return DocxDocument(tmp)
    finally:
        os.unlink(tmp)


def test_render_heading():
    doc = _render_and_check([
        Heading(level=1, runs=[TextRun("Title")]),
    ])
    assert doc.paragraphs[0].style.name == "Heading 1"
    assert doc.paragraphs[0].text == "Title"


def test_render_paragraph():
    doc = _render_and_check([
        Paragraph(runs=[TextRun("Hello")]),
    ])
    assert doc.paragraphs[0].text == "Hello"


def test_render_bold_italic():
    doc = _render_and_check([
        Paragraph(runs=[
            TextRun("bold", bold=True),
            TextRun("normal"),
            TextRun("italic", italic=True),
        ]),
    ])
    p = doc.paragraphs[0]
    assert p.runs[0].bold == True
    assert p.runs[1].bold != True
    assert p.runs[2].italic == True


def test_render_code_block():
    doc = _render_and_check([
        CodeBlock(code="print(1)", language="python"),
    ])
    txt = doc.paragraphs[0].text
    assert "print(1)" in txt


def test_render_list():
    doc = _render_and_check([
        ListBlock(
            ordered=False,
            items=[
                ListItem(elements=[Paragraph(runs=[TextRun("A")])]),
                ListItem(elements=[Paragraph(runs=[TextRun("B")])]),
            ],
        ),
    ])
    # Should have two paragraphs
    paras = [p for p in doc.paragraphs if p.text.strip()]
    assert len(paras) >= 2


def test_render_table():
    doc = _render_and_check([
        Table(
            headers=["Name", "Age"],
            rows=[["Alice", "30"], ["Bob", "25"]],
            align=[None, None],
        ),
    ])
    # The table exists in the document
    tables = doc.tables
    assert len(tables) == 1
    assert tables[0].cell(0, 0).text == "Name"
    assert tables[0].cell(1, 0).text == "Alice"


def test_render_horizontal_rule():
    doc = _render_and_check([HorizontalRule()])
    assert doc.paragraphs[0] is not None


def test_render_empty():
    doc = _render_and_check([])
    assert doc.paragraphs == []
