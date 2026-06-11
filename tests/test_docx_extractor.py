"""Tests for Phase 8 docx reverse extraction features."""
import os
import tempfile
import shutil
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from md2word.extractor.docx_extractor import DocxExtractor
from md2word.writer.md_writer import MdWriter
from md2word.model.document import (
    Heading, Paragraph, CodeBlock, Hyperlink, Image,
    Table, HorizontalRule, PageBreak, Footnote, Comment,
    TextRun,
)


def _make_docx(tmp_dir, paragraphs=None, code_lines=None, page_breaks=None,
               table_data=None, hyperlinks=None):
    """Helper: create a test .docx and return the path."""
    doc = Document()

    if paragraphs:
        for text in paragraphs:
            doc.add_paragraph(text)

    if code_lines:
        for line in code_lines:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Consolas"

    if page_breaks:
        for text in page_breaks:
            doc.add_page_break()
            doc.add_paragraph(text)

    if table_data:
        rows, cols = len(table_data), len(table_data[0])
        table = doc.add_table(rows=rows, cols=cols)
        for r, row in enumerate(table_data):
            for c, cell_text in enumerate(row):
                table.cell(r, c).text = cell_text

    if hyperlinks:
        NS_R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        for url, text in hyperlinks:
            rId_num = 1
            while f"rId{rId_num}" in doc.part.rels:
                rId_num += 1
            rId = f"rId{rId_num}"
            doc.part.rels.add_relationship(
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                url, rId, is_external=True,
            )
            p = OxmlElement("w:p")
            r = OxmlElement("w:r")
            t = OxmlElement("w:t")
            t.text = text
            r.append(t)
            p.append(r)
            hl = OxmlElement("w:hyperlink")
            hl.set(qn("r:id"), rId)
            hr = OxmlElement("w:r")
            ht = OxmlElement("w:t")
            ht.text = text
            hr.append(ht)
            hl.append(hr)
            p.append(hl)
            doc.element.body.append(p)

    path = os.path.join(tmp_dir, "test.docx")
    doc.save(path)
    return path


def _extract(docx_path, tmp_dir):
    """Helper: extract docx and return (document, md_text)."""
    images_dir = os.path.join(tmp_dir, "images")
    os.makedirs(images_dir, exist_ok=True)
    extractor = DocxExtractor(output_dir=images_dir)
    document = extractor.extract(docx_path)
    writer = MdWriter()
    md_text = writer.write(document)
    return document, md_text, extractor


def test_code_block_extraction():
    with tempfile.TemporaryDirectory() as tmp:
        path = _make_docx(tmp, code_lines=["line1", "line2", "line3"])
        doc, md, _ = _extract(path, tmp)
        code_blocks = [e for e in doc.elements if isinstance(e, CodeBlock)]
        assert len(code_blocks) == 1
        assert "line1\nline2\nline3" in code_blocks[0].code
        assert "```" in md


def test_code_block_separated():
    with tempfile.TemporaryDirectory() as tmp:
        doc = Document()
        # First code block
        for line in ["a1", "a2"]:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Consolas"
        # Normal paragraph breaks the code
        doc.add_paragraph("separator")
        # Second code block
        for line in ["b1", "b2"]:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Consolas"

        path = os.path.join(tmp, "test.docx")
        doc.save(path)
        doc_result, md, _ = _extract(path, tmp)
        code_blocks = [e for e in doc_result.elements if isinstance(e, CodeBlock)]
        assert len(code_blocks) == 2
        assert "a1\na2" in code_blocks[0].code
        assert "b1\nb2" in code_blocks[1].code


def test_page_break_extraction():
    with tempfile.TemporaryDirectory() as tmp:
        path = _make_docx(tmp, paragraphs=["Before"], page_breaks=["After"])
        doc, md, _ = _extract(path, tmp)
        page_breaks = [e for e in doc.elements if isinstance(e, PageBreak)]
        assert len(page_breaks) == 1
        assert "\f" in md


def test_table_extraction():
    with tempfile.TemporaryDirectory() as tmp:
        data = [["H1", "H2"], ["a", "b"], ["c", "d"]]
        path = _make_docx(tmp, table_data=data)
        doc, md, _ = _extract(path, tmp)
        tables = [e for e in doc.elements if isinstance(e, Table)]
        assert len(tables) == 1
        assert tables[0].headers == ["H1", "H2"]
        assert tables[0].rows == [["a", "b"], ["c", "d"]]
        assert "| H1 | H2 |" in md
        assert "| a | b |" in md


def test_hyperlink_extraction():
    with tempfile.TemporaryDirectory() as tmp:
        path = _make_docx(tmp, hyperlinks=[("https://example.com", "Example")])
        doc, md, _ = _extract(path, tmp)
        hyperlinks = []
        for e in doc.elements:
            if isinstance(e, Paragraph):
                for run in e.runs:
                    if isinstance(run, Hyperlink):
                        hyperlinks.append(run)
        assert len(hyperlinks) >= 1
        assert hyperlinks[0].url == "https://example.com"
        assert "[Example](https://example.com)" in md


def test_mixed_formatting():
    with tempfile.TemporaryDirectory() as tmp:
        doc = Document()
        p = doc.add_paragraph()
        r1 = p.add_run("bold")
        r1.bold = True
        r2 = p.add_run(" and ")
        r3 = p.add_run("italic")
        r3.italic = True
        path = os.path.join(tmp, "test.docx")
        doc.save(path)
        doc_result, md, _ = _extract(path, tmp)
        assert "**bold**" in md
        assert "*italic*" in md


def test_image_extraction_roundtrip():
    with tempfile.TemporaryDirectory() as tmp:
        # First create a docx with image via forward converter
        md_input = os.path.join(tmp, "input.md")
        with open(md_input, "w", encoding="utf-8") as f:
            f.write("# Test\n\n![alt](image.png){:width=5.77in height=1.60in}\n")

        # Create a dummy image file
        img_path = os.path.join(tmp, "image.png")
        from PIL import Image as PILImage
        img = PILImage.new("RGB", (100, 50), color="red")
        img.save(img_path)

        # Forward convert
        from md2word import MD2Word
        converter = MD2Word()
        docx_path = os.path.join(tmp, "test.docx")
        converter.convert_file(md_input, docx_path)

        # Reverse convert
        images_dir = os.path.join(tmp, "images")
        os.makedirs(images_dir, exist_ok=True)
        extractor = DocxExtractor(output_dir=images_dir)
        doc_result = extractor.extract(docx_path)
        writer = MdWriter()
        md_text = writer.write(doc_result)

        assert len(extractor._saved_images) >= 1
        assert "images/" in md_text


def test_heading_extraction():
    with tempfile.TemporaryDirectory() as tmp:
        doc = Document()
        doc.add_heading("Title", level=1)
        doc.add_heading("Subtitle", level=2)
        doc.add_paragraph("Body")
        path = os.path.join(tmp, "test.docx")
        doc.save(path)
        doc_result, md, _ = _extract(path, tmp)
        headings = [e for e in doc_result.elements if isinstance(e, Heading)]
        assert len(headings) == 2
        assert headings[0].level == 1
        assert headings[1].level == 2
        assert "# Title" in md
        assert "## Subtitle" in md


def test_horizontal_rule_extraction():
    with tempfile.TemporaryDirectory() as tmp:
        doc = Document()
        doc.add_paragraph("Before")
        # Add horizontal rule via paragraph border
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "auto")
        pBdr.append(bottom)
        pPr.append(pBdr)
        doc.add_paragraph("After")
        path = os.path.join(tmp, "test.docx")
        doc.save(path)
        doc_result, md, _ = _extract(path, tmp)
        hrs = [e for e in doc_result.elements if isinstance(e, HorizontalRule)]
        assert len(hrs) == 1
        assert "---" in md
