"""Generate test .docx for Phase 8 features and test reverse conversion."""
from __future__ import annotations
import os, sys, zipfile, shutil
from lxml import etree
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = "output"
os.makedirs(OUTPUT_DIR, exist_ok=True)

doc = Document()

# Normal text
doc.add_paragraph("Normal paragraph with plain text.")

# Monospace code block (consecutive lines)
for line in ["import sys", "", "def main():", "    print('hello world')", "    sys.exit(0)"]:
    p = doc.add_paragraph()
    run = p.add_run(line)
    run.font.name = "Consolas"

doc.add_paragraph("Normal text after code block.")

# Another code block
for line in ["$ ls -la", "$ python test.py"]:
    p = doc.add_paragraph()
    run = p.add_run(line)
    run.font.name = "Courier New"

doc.add_paragraph("End of code.")

# Page break
doc.add_page_break()
doc.add_paragraph("Text on page 2.")
doc.add_page_break()
doc.add_paragraph("Text on page 3.")

# Table
table = doc.add_table(rows=3, cols=3)
table.cell(0, 0).text = "Left"
table.cell(0, 1).text = "Center"
table.cell(0, 2).text = "Right"
table.cell(1, 0).text = "A"
table.cell(1, 1).text = "B"
table.cell(1, 2).text = "C"
table.cell(2, 0).text = "X"
table.cell(2, 1).text = "Y"
table.cell(2, 2).text = "Z"

# Mixed formatting
p = doc.add_paragraph()
run1 = p.add_run("Bold and ")
run1.bold = True
run2 = p.add_run("italic")
run2.italic = True
run3 = p.add_run(" text.")

docx_path = os.path.join(OUTPUT_DIR, "test_phase8_features.docx")
doc.save(docx_path)

# Now inject a hyperlink into the zip
temp_path = docx_path + ".tmp"
with zipfile.ZipFile(docx_path, "r") as zin:
    items = {item.filename: zin.read(item) for item in zin.infolist()}

# Read and modify document.xml
ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
ns_r = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
ns_pkg = "http://schemas.openxmlformats.org/package/2006/relationships"

doc_xml = etree.fromstring(items["word/document.xml"])
body = doc_xml.find(f"{{{ns_w}}}body")
paras = body.findall(f"{{{ns_w}}}p")
last_para = paras[-1]

# Add hyperlink after the last paragraph
hyper_p = etree.SubElement(body, f"{{{ns_w}}}p")
# Add a run before hyperlink
r_before = etree.SubElement(hyper_p, f"{{{ns_w}}}r")
t_before = etree.SubElement(r_before, f"{{{ns_w}}}t")
t_before.text = "Visit "
t_before.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

hyperlink = etree.SubElement(hyper_p, f"{{{ns_w}}}hyperlink")
hyperlink.set(f"{{{ns_r}}}id", "rIdHyperlink")
r_elem = etree.SubElement(hyperlink, f"{{{ns_w}}}r")
rPr = etree.SubElement(r_elem, f"{{{ns_w}}}rPr")
etree.SubElement(rPr, f"{{{ns_w}}}b")
etree.SubElement(rPr, f"{{{ns_w}}}color")
rPr.find(f"{{{ns_w}}}color").set(f"{{{ns_w}}}val", "0563C1")
etree.SubElement(rPr, f"{{{ns_w}}}u")
rPr.find(f"{{{ns_w}}}u").set(f"{{{ns_w}}}val", "single")
t_elem = etree.SubElement(r_elem, f"{{{ns_w}}}t")
t_elem.text = "Example Link"
t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

r_after = etree.SubElement(hyper_p, f"{{{ns_w}}}r")
t_after = etree.SubElement(r_after, f"{{{ns_w}}}t")
t_after.text = " for testing."
t_after.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

items["word/document.xml"] = etree.tostring(doc_xml, xml_declaration=True, encoding="UTF-8", standalone=True)

# Add relationship to rels
rels_xml = etree.fromstring(items["word/_rels/document.xml.rels"])
rel_elem = etree.SubElement(rels_xml, f"{{{ns_pkg}}}Relationship")
rel_elem.set("Id", "rIdHyperlink")
rel_elem.set("Type", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink")
rel_elem.set("Target", "https://example.com")
rel_elem.set("TargetMode", "External")
items["word/_rels/document.xml.rels"] = etree.tostring(rels_xml, xml_declaration=True, encoding="UTF-8", standalone=True)

# Write back
with zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as zout:
    for name, data in items.items():
        zout.writestr(name, data)
shutil.move(temp_path, docx_path)
print(f"Created: {docx_path}")

# ── Test reverse conversion ──
sys.path.insert(0, os.getcwd())
from md2word.extractor.docx_extractor import DocxExtractor
from md2word.writer.md_writer import MdWriter

img_dir = os.path.join(OUTPUT_DIR, "test_phase8_features", "images")
os.makedirs(img_dir, exist_ok=True)
extractor = DocxExtractor(output_dir=img_dir)
document = extractor.extract(docx_path)
writer = MdWriter()
md_text = writer.write(document)

out_md = os.path.join(OUTPUT_DIR, "test_phase8_features", "output.md")
with open(out_md, "w", encoding="utf-8") as f:
    f.write(md_text)
print(f"Written: {out_md}")
print(f"Images extracted: {len(extractor._saved_images)}")
print()
print("=== OUTPUT (first 1500 chars) ===")
print(md_text[:1500])
print("=== END ===")
