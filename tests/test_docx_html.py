import os
import subprocess
import sys
import tempfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image as PILImage

from md2word.html import render_docx_to_html, render_metadata_to_html
from md2word.meta import DocxMetaExtractor


def _set_numbering(paragraph, num_id="5", level="0"):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), level)
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), num_id)
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)


def _add_simple_math(run, text):
    math = OxmlElement("m:oMath")
    math_run = OxmlElement("m:r")
    math_text = OxmlElement("m:t")
    math_text.text = text
    math_run.append(math_text)
    math.append(math_run)
    run._r.append(math)


def test_docx_html_renderer_uses_metadata_semantics():
    with tempfile.TemporaryDirectory() as tmp:
        image_path = os.path.join(tmp, "image.png")
        PILImage.new("RGB", (18, 12), color="orange").save(image_path)

        doc = Document()
        doc.sections[0].header.paragraphs[0].text = "Header text"
        doc.sections[0].footer.paragraphs[0].text = "Footer text"
        doc.add_heading("HTML Coverage", level=1)

        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run("Styled")
        run.bold = True
        run.underline = True
        run.font.name = "Arial"
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x12, 0x34, 0x56)
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

        list_paragraph = doc.add_paragraph("Numbered")
        _set_numbering(list_paragraph)

        math_paragraph = doc.add_paragraph("Math: ")
        _add_simple_math(math_paragraph.add_run(), "x+1")

        doc.add_picture(image_path, width=Inches(1))

        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Merged"
        table.cell(0, 0).merge(table.cell(0, 1))
        table.cell(1, 0).text = "A"
        table.cell(1, 1).text = "B"

        source_docx = os.path.join(tmp, "source.docx")
        doc.save(source_docx)

        html_path = os.path.join(tmp, "source.html")
        render_docx_to_html(source_docx, html_path)

        html = open(html_path, "r", encoding="utf-8").read()
        assert "MathJax" in html
        assert "<h1" in html and "HTML Coverage" in html
        assert "Header text" in html
        assert "Footer text" in html
        assert "text-align:center" in html
        assert "color:#123456" in html
        assert "background-color:#ffff00" in html
        assert "data:image/png;base64," in html
        assert 'colspan="2"' in html
        assert "<ol" in html and "Numbered" in html
        assert "\\(x+1\\)" in html

        meta_dir = os.path.join(tmp, "meta")
        DocxMetaExtractor().extract(source_docx, meta_dir)
        html_from_meta = os.path.join(tmp, "source-from-meta.html")
        render_metadata_to_html(meta_dir, html_from_meta)
        assert os.path.exists(html_from_meta)
        assert "HTML Coverage" in open(html_from_meta, "r", encoding="utf-8").read()


def test_to_html_cli_writes_output_file():
    with tempfile.TemporaryDirectory() as tmp:
        doc = Document()
        doc.add_paragraph("CLI HTML")
        source_docx = os.path.join(tmp, "cli.docx")
        html_path = os.path.join(tmp, "cli.html")
        doc.save(source_docx)

        result = subprocess.run(
            [
                sys.executable,
                "-m",
                "md2word.cli.main",
                source_docx,
                "--to-html",
                "-o",
                html_path,
            ],
            cwd=os.getcwd(),
            capture_output=True,
            text=True,
            check=False,
        )

        assert result.returncode == 0, result.stderr
        assert os.path.exists(html_path)
        assert "CLI HTML" in open(html_path, "r", encoding="utf-8").read()
