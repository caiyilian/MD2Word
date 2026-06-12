from __future__ import annotations
import os
from typing import List, Optional

from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

from md2word.model.document import (
    TextRun, Image, Heading, Paragraph, CodeBlock, Hyperlink,
    ListBlock, ListItem, Table, HorizontalRule, Formula, PageBreak,
    Footnote, Comment, Blockquote,
    Document, InlineElement, BlockElement,
)
from md2word.renderer.formula_converter import latex_to_omml
from md2word.utils.unit_converter import parse_size
from md2word.renderer.styles import load_style_config, parse_color


class DocxRenderer:
    def __init__(self, font_name: str = "等线", font_size: int = 12,
                 base_dir: str = "", style_config: Optional[dict] = None):
        self.font_name = font_name
        self.font_size = font_size
        self.base_dir = base_dir
        self.style = style_config or {}

    def render(self, document: Document, output_path: str):
        doc = DocxDocument()
        self._apply_style_config(doc)
        self._set_default_style(doc)

        # Render headers and footers from metadata
        if document.headers or document.footers:
            self._render_headers_footers(doc, document)

        # Collect footnotes and comments from document
        footnotes = [e for e in document.elements if isinstance(e, Footnote)]
        comments = [e for e in document.elements if isinstance(e, Comment)]

        # Create footnotes part if needed
        if footnotes:
            self._create_footnotes_part(doc, footnotes)

        # Create comments part if needed
        if comments:
            self._create_comments_part(doc, comments)

        # Render non-footnote, non-comment elements
        for element in document.elements:
            if not isinstance(element, (Footnote, Comment)):
                self._render_element(doc, element)

        # Post-process: insert footnote references and comment ranges
        if footnotes:
            self._insert_footnote_references(doc)
        if comments:
            self._insert_comment_ranges(doc, comments)

        # Insert TOC field if requested
        if document.metadata.get("toc"):
            self._insert_toc_field(doc)

        # Insert bookmarks if requested
        if document.metadata.get("bookmarks"):
            self._insert_bookmarks(doc)

        doc.save(output_path)
        return output_path

    def _set_default_style(self, doc: DocxDocument):
        style = doc.styles["Normal"]
        font = style.font
        font.name = self.font_name
        font.size = Pt(self.font_size)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.15

    def _apply_style_config(self, doc: DocxDocument):
        if not self.style:
            return

        # Body font
        body_cfg = self.style.get("body", {})
        if body_cfg.get("font"):
            self.font_name = body_cfg["font"]
        if body_cfg.get("size"):
            self.font_size = body_cfg["size"]

        ns = doc.styles["Normal"]
        ns.font.name = self.font_name
        ns.font.size = Pt(self.font_size)
        color = parse_color(body_cfg.get("color"))
        if color:
            ns.font.color.rgb = color
        if body_cfg.get("spacing"):
            ns.paragraph_format.line_spacing = body_cfg["spacing"]

        # Page margins
        page_cfg = self.style.get("page", {})
        sections = doc.sections
        if sections:
            sec = sections[0]
            for attr, key in [("top", "margin_top"), ("bottom", "margin_bottom"),
                              ("left", "margin_left"), ("right", "margin_right")]:
                val = page_cfg.get(key)
                if val is not None:
                    setattr(sec, attr, Inches(val))

        # Column layout
        columns = self.style.get("columns", 1)
        if columns > 1:
            sec = doc.sections[0]
            sectPr = sec._sectPr
            ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            # Remove existing cols element
            existing_cols = sectPr.find("{%s}cols" % ns_w)
            if existing_cols is not None:
                sectPr.remove(existing_cols)
            # Create new cols element
            cols_xml = '<w:cols xmlns:w="%s" w:num="%d" w:space="720"/>' % (ns_w, columns)
            cols = etree.fromstring(cols_xml.encode())
            sectPr.append(cols)

    def _render_element(self, doc: DocxDocument, element: BlockElement):
        if isinstance(element, Heading):
            self._render_heading(doc, element)
        elif isinstance(element, Paragraph):
            self._render_paragraph(doc, element)
        elif isinstance(element, CodeBlock):
            self._render_code_block(doc, element)
        elif isinstance(element, ListBlock):
            self._render_list(doc, element, indent_level=0)
        elif isinstance(element, Table):
            self._render_table(doc, element)
        elif isinstance(element, Image):
            self._render_image(doc, element)
        elif isinstance(element, HorizontalRule):
            self._render_horizontal_rule(doc)
        elif isinstance(element, Formula):
            self._render_formula(doc, element)
        elif isinstance(element, PageBreak):
            self._render_page_break(doc)
        elif isinstance(element, Blockquote):
            self._render_blockquote(doc, element)
        # Footnote and Comment are handled in render() post-processing

    def _render_heading(self, doc: DocxDocument, heading: Heading):
        style_name = f"Heading {heading.level}"
        p = doc.add_paragraph(style=style_name)
        # Apply heading color from style config
        h_key = f"h{heading.level}"
        h_cfg = self.style.get("headings", {}).get(h_key, {})
        h_color = parse_color(h_cfg.get("color"))
        if h_color:
            for run in p.runs:
                run.font.color.rgb = h_color
        self._apply_runs(p, heading.runs)

    def _render_paragraph(self, doc: DocxDocument, paragraph: Paragraph):
        p = doc.add_paragraph()
        self._apply_runs(p, paragraph.runs)

    def _render_code_block(self, doc: DocxDocument, code_block: CodeBlock):
        code_cfg = self.style.get("code", {})

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.3)

        bg = code_cfg.get("bg_color", "F2F2F2")
        pPr = p._p.get_or_add_pPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), bg)
        shading.set(qn("w:val"), "clear")
        pPr.append(shading)

        run = p.add_run(code_block.code)
        run.font.name = code_cfg.get("font", "Consolas")
        run.font.size = Pt(code_cfg.get("size", 9))
        color = parse_color(code_cfg.get("color", "333333"))
        if color:
            run.font.color.rgb = color
        else:
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    def _render_image(self, doc: DocxDocument, image: Image):
        p = doc.add_paragraph()
        self._apply_image_alignment(p, image.align)
        if not self._add_image_to_paragraph(p, image):
            run = p.add_run(f"[图片未找到: {image.src}]")
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            run.font.size = Pt(9)

    def _render_table(self, doc: DocxDocument, table: Table):
        if not table.headers and not table.rows:
            return

        num_cols = max(len(table.headers), max((len(r) for r in table.rows), default=0))
        if num_cols == 0:
            return

        t = doc.add_table(rows=1 + len(table.rows), cols=num_cols)
        t.style = "Table Grid"
        t.autofit = True

        # Headers
        if table.headers:
            for i, header in enumerate(table.headers):
                if i >= num_cols:
                    break
                cell = t.cell(0, i)
                cell.text = ""
                run = cell.paragraphs[0].add_run(header)
                run.bold = True
                run.font.name = self.font_name
                run.font.size = Pt(self.font_size)
                # Header shading
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "E8E8E8")
                shading.set(qn("w:val"), "clear")
                cell._tc.get_or_add_tcPr().append(shading)

        # Rows
        for row_idx, row_data in enumerate(table.rows):
            for col_idx, cell_text in enumerate(row_data):
                if col_idx >= num_cols:
                    break
                cell = t.cell(1 + row_idx, col_idx)
                cell.text = ""
                run = cell.paragraphs[0].add_run(cell_text)
                run.font.name = self.font_name
                run.font.size = Pt(self.font_size)

        # Column alignment
        for i, align_val in enumerate(table.align):
            if i >= num_cols:
                break
            if align_val == "left":
                alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif align_val == "center":
                alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align_val == "right":
                alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                continue
            # Set alignment for header
            if table.headers:
                t.cell(0, i).paragraphs[0].alignment = alignment
            # Set alignment for body cells
            for row_idx in range(len(table.rows)):
                t.cell(1 + row_idx, i).paragraphs[0].alignment = alignment

        doc.add_paragraph()  # spacing after table

    def _render_list(self, doc: DocxDocument, list_block: ListBlock,
                     indent_level: int = 0):
        for idx, item in enumerate(list_block.items, start=1):
            self._render_list_item(doc, item, list_block.ordered, idx, indent_level)

    def _render_list_item(self, doc: DocxDocument, item: ListItem,
                          ordered: bool, idx: int, indent_level: int):
        for element in item.elements:
            if isinstance(element, Paragraph):
                # Handle task list checkboxes
                if item.checked is not None:
                    checkbox = "\u2611 " if item.checked else "\u2610 "  # ☑ or ☐
                    style_name = "List Bullet"
                    if indent_level > 0:
                        style_name += f" {indent_level + 1}"
                    try:
                        p = doc.add_paragraph(style=style_name)
                    except Exception:
                        p = doc.add_paragraph()
                        indent = Inches(0.5 + indent_level * 0.4)
                        p.paragraph_format.left_indent = indent
                    # Add checkbox character
                    checkbox_run = p.add_run(checkbox)
                    checkbox_run.font.name = self.font_name
                    checkbox_run.font.size = Pt(self.font_size)
                    self._apply_runs(p, element.runs)
                elif ordered:
                    style_name = f"List Number"
                    if indent_level > 0:
                        style_name += f" {indent_level + 1}"
                    try:
                        p = doc.add_paragraph(style=style_name)
                    except Exception:
                        p = doc.add_paragraph()
                        indent = Inches(0.5 + indent_level * 0.4)
                        p.paragraph_format.left_indent = indent
                        prefix = f"{idx}. "
                        prefix_run = p.add_run(prefix)
                        prefix_run.font.name = self.font_name
                        prefix_run.font.size = Pt(self.font_size)
                    self._apply_runs(p, element.runs)
                else:
                    style_name = f"List Bullet"
                    if indent_level > 0:
                        style_name += f" {indent_level + 1}"
                    try:
                        p = doc.add_paragraph(style=style_name)
                    except Exception:
                        p = doc.add_paragraph()
                        indent = Inches(0.5 + indent_level * 0.4)
                        p.paragraph_format.left_indent = indent
                        prefix = "\u2022 "
                        prefix_run = p.add_run(prefix)
                        prefix_run.font.name = self.font_name
                        prefix_run.font.size = Pt(self.font_size)
                    self._apply_runs(p, element.runs)

            elif isinstance(element, Image):
                if ordered:
                    style_name = f"List Number"
                else:
                    style_name = f"List Bullet"
                if indent_level > 0:
                    style_name += f" {indent_level + 1}"

                try:
                    p = doc.add_paragraph(style=style_name)
                except Exception:
                    p = doc.add_paragraph()
                    indent = Inches(0.5 + indent_level * 0.4)
                    p.paragraph_format.left_indent = indent
                    prefix = f"{idx}. " if ordered else "\u2022 "
                    prefix_run = p.add_run(prefix)
                    prefix_run.font.name = self.font_name
                    prefix_run.font.size = Pt(self.font_size)

                self._add_image_to_paragraph(p, element)
            elif isinstance(element, CodeBlock):
                self._render_code_block(doc, element)
            elif isinstance(element, Formula):
                self._render_formula(doc, element)
            elif isinstance(element, ListBlock):
                self._render_list(doc, element, indent_level + 1)

    def _render_horizontal_rule(self, doc: DocxDocument):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)

        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "999999")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _render_blockquote(self, doc: DocxDocument, blockquote: Blockquote):
        # Build text from runs
        text_parts = []
        for run in blockquote.runs:
            if isinstance(run, TextRun):
                text_parts.append(run.text)
            elif isinstance(run, Hyperlink):
                text_parts.append(run.url)
        text = "".join(text_parts)

        lines = text.split("\n")

        for i, line in enumerate(lines):
            if not line.strip() and i > 0:
                continue  # Skip empty lines between blockquote lines
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5 * blockquote.level)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

            # Add left border (blue-gray color)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"), "single")
            left.set(qn("w:sz"), "12")
            left.set(qn("w:space"), "4")
            left.set(qn("w:color"), "4472C4")
            pBdr.append(left)
            pPr.append(pBdr)

            # Add shading (light gray background)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "F2F2F2")
            shading.set(qn("w:val"), "clear")
            pPr.append(shading)

            # Add the text
            run = p.add_run(line.strip())
            run.font.name = self.font_name
            run.font.size = Pt(self.font_size)

    # ---- helpers ----

    def _render_formula(self, doc: DocxDocument, formula: Formula):
        p = doc.add_paragraph()
        if formula.display:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._insert_omml(p, formula.latex)
        if formula.numbering is not None:
            run = p.add_run(f"  ({formula.numbering})")
            run.font.size = Pt(self.font_size)
            run.font.name = self.font_name

    def _insert_omml(self, paragraph, latex: str):
        omml_str = latex_to_omml(latex)
        if omml_str is None:
            paragraph.add_run(f"[公式解析失败: {latex}]")
            return
        from lxml import etree
        omml_elem = etree.fromstring(omml_str.encode("utf-8"))
        paragraph._p.append(omml_elem)

    def _apply_runs(self, paragraph, runs: List[InlineElement]):
        for run_data in runs:
            if isinstance(run_data, Image):
                self._add_image_to_paragraph(paragraph, run_data)
            elif isinstance(run_data, Formula):
                self._insert_omml(paragraph, run_data.latex)
            elif isinstance(run_data, Hyperlink):
                self._render_hyperlink(paragraph, run_data)
            elif isinstance(run_data, TextRun):
                if not run_data.text:
                    continue
                r = paragraph.add_run(run_data.text)
                r.bold = run_data.bold
                r.italic = run_data.italic
                r.underline = run_data.underline
                r.font.strike = run_data.strikethrough
                r.font.superscript = run_data.superscript
                r.font.subscript = run_data.subscript

                if run_data.code:
                    r.font.name = "Consolas"
                    r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                else:
                    r.font.name = run_data.font_name or self.font_name
                    r.font.size = Pt(run_data.font_size or self.font_size)

    def _add_image_to_paragraph(self, paragraph, image: Image):
        img_path = self._resolve_image_path(image.src)
        if img_path is None:
            return False

        try:
            pw_twips = paragraph.part.document.element.body.sectPr.xpath(
                "./w:pgSz/@w:w"
            )
            if pw_twips:
                pw = Emu(int(pw_twips[0]) * 635)
            else:
                pw = None
        except Exception:
            pw = None

        width = parse_size(image.width, pw)
        height = parse_size(image.height)

        run = paragraph.add_run()
        run.add_picture(img_path, width=width, height=height)
        return True

    def _resolve_image_path(self, src: str) -> Optional[str]:
        if os.path.isabs(src):
            return src if os.path.exists(src) else None
        if self.base_dir:
            resolved = os.path.join(self.base_dir, src)
            if os.path.exists(resolved):
                return resolved
        return src if os.path.exists(src) else None

    def _apply_image_alignment(self, paragraph, align: Optional[str]):
        if align == "center":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif align == "left":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # ---- Footnotes ----

    def _create_footnotes_part(self, doc: DocxDocument, footnotes: List[Footnote]):
        """Create word/footnotes.xml with footnote text."""
        from lxml import etree

        NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        NS_R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

        # Build footnotes XML
        footnotes_xml = etree.Element(f"{{{NS_W}}}footnotes",
                                      nsmap={"w": NS_W, "r": NS_R})

        # Separator footnote (id=-1)
        sep_fn = etree.SubElement(footnotes_xml, f"{{{NS_W}}}footnote")
        sep_fn.set(f"{{{NS_W}}}type", "separator")
        sep_fn.set(f"{{{NS_W}}}id", "-1")
        sep_p = etree.SubElement(sep_fn, f"{{{NS_W}}}p")
        sep_pPr = etree.SubElement(sep_p, f"{{{NS_W}}}pPr")
        sep_spacing = etree.SubElement(sep_pPr, f"{{{NS_W}}}spacing")
        sep_spacing.set(f"{{{NS_W}}}after", "0")
        sep_spacing.set(f"{{{NS_W}}}line", "240")
        sep_spacing.set(f"{{{NS_W}}}lineRule", "auto")
        sep_r = etree.SubElement(sep_p, f"{{{NS_W}}}r")
        etree.SubElement(sep_r, f"{{{NS_W}}}separator")

        # ContinuationSeparator footnote (id=0)
        cont_fn = etree.SubElement(footnotes_xml, f"{{{NS_W}}}footnote")
        cont_fn.set(f"{{{NS_W}}}type", "continuationSeparator")
        cont_fn.set(f"{{{NS_W}}}id", "0")
        cont_p = etree.SubElement(cont_fn, f"{{{NS_W}}}p")
        cont_pPr = etree.SubElement(cont_p, f"{{{NS_W}}}pPr")
        cont_spacing = etree.SubElement(cont_pPr, f"{{{NS_W}}}spacing")
        cont_spacing.set(f"{{{NS_W}}}after", "0")
        cont_spacing.set(f"{{{NS_W}}}line", "240")
        cont_spacing.set(f"{{{NS_W}}}lineRule", "auto")
        cont_r = etree.SubElement(cont_p, f"{{{NS_W}}}r")
        etree.SubElement(cont_r, f"{{{NS_W}}}continuationSeparator")

        # Add actual footnotes
        for idx, fn in enumerate(footnotes, start=2):
            fn_elem = etree.SubElement(footnotes_xml, f"{{{NS_W}}}footnote")
            fn_elem.set(f"{{{NS_W}}}id", str(idx))
            p = etree.SubElement(fn_elem, f"{{{NS_W}}}p")
            pPr = etree.SubElement(p, f"{{{NS_W}}}pPr")
            pStyle = etree.SubElement(pPr, f"{{{NS_W}}}pStyle")
            pStyle.set(f"{{{NS_W}}}val", "FootnoteText")
            # Footnote reference marker
            r_ref = etree.SubElement(p, f"{{{NS_W}}}r")
            rPr_ref = etree.SubElement(r_ref, f"{{{NS_W}}}rPr")
            rStyle = etree.SubElement(rPr_ref, f"{{{NS_W}}}rStyle")
            rStyle.set(f"{{{NS_W}}}val", "FootnoteReference")
            vertAlign = etree.SubElement(rPr_ref, f"{{{NS_W}}}vertAlign")
            vertAlign.set(f"{{{NS_W}}}val", "superscript")
            etree.SubElement(r_ref, f"{{{NS_W}}}footnoteRef")
            # Footnote text
            r_text = etree.SubElement(p, f"{{{NS_W}}}r")
            t = etree.SubElement(r_text, f"{{{NS_W}}}t")
            t.text = " " + fn.text
            t.set(f"{{{NS_W}}}space", "preserve")

        # Save to part
        blob = etree.tostring(footnotes_xml, xml_declaration=True,
                              encoding="UTF-8", standalone=True)
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI
        part = Part(
            PackURI("/word/footnotes.xml"),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml",
            blob, doc.part.package,
        )
        doc.part.relate_to(
            part,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes",
        )

        # Store footnote id mapping for reference insertion
        self._footnote_id_map = {fn.footnote_id: idx for idx, fn in enumerate(footnotes, start=2)}

    def _insert_footnote_references(self, doc: DocxDocument):
        """Replace [^id] text in runs with footnote references."""
        import re
        NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        FN_REF_RE = re.compile(r'\[\^(\w+)\]')

        for para in doc.paragraphs:
            # Join all run text to find footnote references
            full_text = "".join(r.text for r in para.runs)
            if not FN_REF_RE.search(full_text):
                continue

            # Clear all runs
            for run in para.runs:
                run.text = ""

            # Rebuild text with footnote references
            parts = FN_REF_RE.split(full_text)
            # parts: ['text', 'id1', 'text', 'id2', 'text', ...]

            # Use the first run for text, create new runs for the rest
            first_run = para.runs[0] if para.runs else para.add_run()
            run_idx = 0

            for i, part in enumerate(parts):
                if i % 2 == 0:
                    # Regular text
                    if part:
                        if run_idx == 0:
                            first_run.text = part
                        else:
                            new_run = para.add_run(part)
                            new_run.font.name = first_run.font.name
                            new_run.font.size = first_run.font.size
                        run_idx += 1
                else:
                    # Footnote reference id
                    fn_id = part
                    if fn_id in self._footnote_id_map:
                        fn_num = self._footnote_id_map[fn_id]
                        # Insert footnote reference
                        r_elem = OxmlElement("w:r")
                        rPr = OxmlElement("w:rPr")
                        vertAlign = OxmlElement("w:vertAlign")
                        vertAlign.set(qn("w:val"), "superscript")
                        rPr.append(vertAlign)
                        r_elem.append(rPr)
                        fn_ref = OxmlElement("w:footnoteReference")
                        fn_ref.set(qn("w:id"), str(fn_num))
                        r_elem.append(fn_ref)
                        para._p.append(r_elem)

    # ---- Comments ----

    def _create_comments_part(self, doc: DocxDocument, comments: List[Comment]):
        """Create word/comments.xml with comment text."""
        from lxml import etree

        NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        NS_R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

        comments_xml = etree.Element(f"{{{NS_W}}}comments",
                                     nsmap={"w": NS_W, "r": NS_R})

        self._comment_id_map = {}
        for idx, cm in enumerate(comments):
            cm_id = str(idx)
            self._comment_id_map[cm.text] = cm_id

            cm_elem = etree.SubElement(comments_xml, f"{{{NS_W}}}comment")
            cm_elem.set(f"{{{NS_W}}}id", cm_id)
            cm_elem.set(f"{{{NS_W}}}author", cm.author or "Anonymous")
            cm_elem.set(f"{{{NS_W}}}date", cm.date or "2026-01-01T00:00:00Z")
            cm_elem.set(f"{{{NS_W}}}initials", (cm.author or "A")[0])

            p = etree.SubElement(cm_elem, f"{{{NS_W}}}p")
            r = etree.SubElement(p, f"{{{NS_W}}}r")
            t = etree.SubElement(r, f"{{{NS_W}}}t")
            t.text = cm.text
            t.set(f"{{{NS_W}}}space", "preserve")

        # Save to part
        blob = etree.tostring(comments_xml, xml_declaration=True,
                              encoding="UTF-8", standalone=True)
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI
        part = Part(
            PackURI("/word/comments.xml"),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml",
            blob, doc.part.package,
        )
        doc.part.relate_to(
            part,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments",
        )

    def _insert_comment_ranges(self, doc: DocxDocument, comments: List[Comment]):
        """Insert comment range markers around target text in paragraphs."""
        if not comments:
            return

        for idx, cm in enumerate(comments):
            cm_id = self._comment_id_map.get(cm.text, str(idx))

            if not cm.target:
                # No target, wrap whole paragraph
                paras = [p for p in doc.paragraphs
                         if p.text.strip()
                         and not (p.style and p.style.name.startswith("Heading"))]
                if idx < len(paras):
                    self._wrap_paragraph_with_comment(paras[idx], cm_id)
                continue

            # Find paragraph containing target text (prefer last match)
            target_para = None
            for para in doc.paragraphs:
                if para.style and para.style.name.startswith("Heading"):
                    continue
                full_text = "".join(r.text for r in para.runs)
                if cm.target in full_text:
                    target_para = para  # Keep last match

            if target_para is None:
                continue

            # Try to wrap just the target text precisely
            if not self._wrap_text_with_comment(target_para, cm.target, cm_id):
                # Fallback: wrap whole paragraph
                self._wrap_paragraph_with_comment(target_para, cm_id)

    def _wrap_text_with_comment(self, paragraph, target: str, cm_id: str) -> bool:
        """Wrap specific target text with comment markers. Returns True if successful."""
        runs = paragraph.runs
        if not runs:
            return False

        # Find which run(s) contain the target text
        for run_idx, run in enumerate(runs):
            if target not in run.text:
                continue

            # Found the run containing the target
            text = run.text
            pos = text.find(target)
            before = text[:pos]
            after = text[pos + len(target):]

            # Clear the original run
            run.text = ""

            # Get run formatting for new runs
            font_name = run.font.name
            font_size = run.font.size

            # Create run for text before target (inserted first)
            if before:
                r_before = OxmlElement("w:r")
                rPr = OxmlElement("w:rPr")
                if font_name:
                    rFonts = OxmlElement("w:rFonts")
                    rFonts.set(qn("w:ascii"), font_name)
                    rFonts.set(qn("w:hAnsi"), font_name)
                    rPr.append(rFonts)
                if font_size:
                    sz = OxmlElement("w:sz")
                    sz.set(qn("w:val"), str(int(font_size.pt * 2)))
                    rPr.append(sz)
                r_before.append(rPr)
                t_before = OxmlElement("w:t")
                t_before.text = before
                t_before.set(qn("w:space"), "preserve")
                r_before.append(t_before)
                run._r.addprevious(r_before)

            # Insert commentRangeStart BEFORE target run
            range_start = OxmlElement("w:commentRangeStart")
            range_start.set(qn("w:id"), cm_id)
            run._r.addprevious(range_start)

            # Create run for the target text
            r_target = OxmlElement("w:r")
            rPr = OxmlElement("w:rPr")
            if font_name:
                rFonts = OxmlElement("w:rFonts")
                rFonts.set(qn("w:ascii"), font_name)
                rFonts.set(qn("w:hAnsi"), font_name)
                rPr.append(rFonts)
            if font_size:
                sz = OxmlElement("w:sz")
                sz.set(qn("w:val"), str(int(font_size.pt * 2)))
                rPr.append(sz)
            r_target.append(rPr)
            t_target = OxmlElement("w:t")
            t_target.text = target
            t_target.set(qn("w:space"), "preserve")
            r_target.append(t_target)
            run._r.addprevious(r_target)

            # Insert commentRangeEnd AFTER target run
            range_end = OxmlElement("w:commentRangeEnd")
            range_end.set(qn("w:id"), cm_id)
            run._r.addprevious(range_end)

            # Create run for text after target
            if after:
                r_after = OxmlElement("w:r")
                rPr = OxmlElement("w:rPr")
                if font_name:
                    rFonts = OxmlElement("w:rFonts")
                    rFonts.set(qn("w:ascii"), font_name)
                    rFonts.set(qn("w:hAnsi"), font_name)
                    rPr.append(rFonts)
                if font_size:
                    sz = OxmlElement("w:sz")
                    sz.set(qn("w:val"), str(int(font_size.pt * 2)))
                    rPr.append(sz)
                r_after.append(rPr)
                t_after = OxmlElement("w:t")
                t_after.text = after
                t_after.set(qn("w:space"), "preserve")
                r_after.append(t_after)
                run._r.addprevious(r_after)

            # Add comment reference at end of paragraph
            r_ref = OxmlElement("w:r")
            rPr = OxmlElement("w:rPr")
            rStyle = OxmlElement("w:rStyle")
            rStyle.set(qn("w:val"), "CommentReference")
            rPr.append(rStyle)
            r_ref.append(rPr)
            cm_ref = OxmlElement("w:commentReference")
            cm_ref.set(qn("w:id"), cm_id)
            r_ref.append(cm_ref)
            paragraph._p.append(r_ref)

            return True

        return False

    def _wrap_paragraph_with_comment(self, paragraph, cm_id: str):
        """Wrap entire paragraph with comment markers."""
        range_start = OxmlElement("w:commentRangeStart")
        range_start.set(qn("w:id"), cm_id)
        paragraph._p.insert(0, range_start)

        range_end = OxmlElement("w:commentRangeEnd")
        range_end.set(qn("w:id"), cm_id)
        paragraph._p.append(range_end)

        r_elem = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rStyle = OxmlElement("w:rStyle")
        rStyle.set(qn("w:val"), "CommentReference")
        rPr.append(rStyle)
        r_elem.append(rPr)
        cm_ref = OxmlElement("w:commentReference")
        cm_ref.set(qn("w:id"), cm_id)
        r_elem.append(cm_ref)
        paragraph._p.append(r_elem)

    def _render_headers_footers(self, doc: DocxDocument, document: Document):
        """Render headers and footers from document metadata."""
        section = doc.sections[0] if doc.sections else None
        if not section:
            return

        # Check if page numbers are enabled in metadata
        page_numbers = document.metadata.get("page_numbers", False)

        # Render headers
        if document.headers:
            header = section.header
            header.is_linked_to_previous = False
            for i, header_text in enumerate(document.headers):
                if i < len(header.paragraphs):
                    p = header.paragraphs[i]
                else:
                    p = header.add_paragraph()
                p.text = header_text
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Render footers
        if document.footers or page_numbers:
            footer = section.footer
            footer.is_linked_to_previous = False

            if document.footers:
                for i, footer_text in enumerate(document.footers):
                    if i < len(footer.paragraphs):
                        p = footer.paragraphs[i]
                    else:
                        p = footer.add_paragraph()
                    p.text = footer_text
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Add page number after footer text if enabled
                    if page_numbers:
                        self._add_page_number_field(p)
            elif page_numbers:
                # No footer text, just page numbers
                if footer.paragraphs:
                    p = footer.paragraphs[0]
                else:
                    p = footer.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._add_page_number_field(p)

    def _add_page_number_field(self, paragraph):
        """Add page number field to a paragraph."""
        # Create run for page number
        run = paragraph.add_run()
        run.font.size = Pt(9)

        # Create the PAGE field
        fldChar_begin = OxmlElement("w:fldChar")
        fldChar_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fldChar_begin)

        # Create the instruction text
        run2 = paragraph.add_run()
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = " PAGE "
        run2._r.append(instrText)

        # Create the end field
        run3 = paragraph.add_run()
        fldChar_end = OxmlElement("w:fldChar")
        fldChar_end.set(qn("w:fldCharType"), "end")
        run3._r.append(fldChar_end)

    def _insert_toc_field(self, doc: DocxDocument):
        """Insert Table of Contents field at the beginning of the document."""
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)

        # Add TOC field
        r1 = OxmlElement("w:r")
        fldChar_begin = OxmlElement("w:fldChar")
        fldChar_begin.set(qn("w:fldCharType"), "begin")
        r1.append(fldChar_begin)
        p._p.append(r1)

        r2 = OxmlElement("w:r")
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
        r2.append(instrText)
        p._p.append(r2)

        r3 = OxmlElement("w:r")
        fldChar_separate = OxmlElement("w:fldChar")
        fldChar_separate.set(qn("w:fldCharType"), "separate")
        r3.append(fldChar_separate)
        p._p.append(r3)

        r4 = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = "Right-click to update field."
        r4.append(t)
        p._p.append(r4)

        r5 = OxmlElement("w:r")
        fldChar_end = OxmlElement("w:fldChar")
        fldChar_end.set(qn("w:fldCharType"), "end")
        r5.append(fldChar_end)
        p._p.append(r5)

        # Move TOC to the beginning of the document
        p._p.getparent().remove(p._p)
        doc.element.body.insert(0, p._p)

    def _insert_bookmarks(self, doc: DocxDocument):
        """Insert bookmarks for headings."""
        for i, para in enumerate(doc.paragraphs):
            if para.style.name.startswith("Heading"):
                # Add bookmark start
                bookmark_start = OxmlElement("w:bookmarkStart")
                bookmark_start.set(qn("w:id"), str(i))
                bookmark_start.set(qn("w:name"), f"heading_{i}")
                para._p.insert(0, bookmark_start)

                # Add bookmark end
                bookmark_end = OxmlElement("w:bookmarkEnd")
                bookmark_end.set(qn("w:id"), str(i))
                para._p.append(bookmark_end)

    def _render_page_break(self, doc: DocxDocument):
        p = doc.add_paragraph()
        run = p.add_run()
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        run._r.append(br)

    def _render_hyperlink(self, paragraph, hyperlink: Hyperlink):
        """Render a hyperlink as w:hyperlink element."""
        NS_R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

        # Add relationship
        rId_num = 1
        while f"rId{rId_num}" in paragraph.part.rels:
            rId_num += 1
        rId = f"rId{rId_num}"
        paragraph.part.rels.add_relationship(
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            hyperlink.url, rId, is_external=True,
        )

        # Create hyperlink element
        hl_elem = OxmlElement("w:hyperlink")
        hl_elem.set(qn("r:id"), rId)

        # Add runs inside hyperlink
        for run_data in hyperlink.runs:
            r = OxmlElement("w:r")
            rPr = OxmlElement("w:rPr")
            color = OxmlElement("w:color")
            color.set(qn("w:val"), "0563C1")
            rPr.append(color)
            u = OxmlElement("w:u")
            u.set(qn("w:val"), "single")
            rPr.append(u)
            if run_data.bold:
                b = OxmlElement("w:b")
                rPr.append(b)
            if run_data.italic:
                i = OxmlElement("w:i")
                rPr.append(i)
            r.append(rPr)
            t = OxmlElement("w:t")
            t.text = run_data.text
            r.append(t)
            hl_elem.append(r)

        paragraph._p.append(hl_elem)
