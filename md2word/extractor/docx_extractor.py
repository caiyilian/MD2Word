from __future__ import annotations
import os
from typing import Dict, List, Optional, Set

from docx import Document as DocxDocument
from docx.oxml.ns import qn
from lxml import etree

from md2word.model.document import (
    TextRun, Image, Heading, Paragraph, CodeBlock, Hyperlink,
    ListBlock, ListItem, Table, HorizontalRule, Formula, PageBreak,
    Footnote, Comment,
    Document, InlineElement, BlockElement,
)


_MONO_FONTS: Set[str] = {
    "consolas", "courier new", "courier", "monospace",
    "source code pro", "fira code", "menlo", "monaco",
    "droid sans mono", "dejavu sans mono", "liberation mono",
}


def _is_mono_font(name: Optional[str]) -> bool:
    return name is not None and name.lower().strip() in _MONO_FONTS


def _alignment_to_str(align) -> Optional[str]:
    if align is None:
        return None
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    if align == WD_ALIGN_PARAGRAPH.LEFT:
        return "left"
    if align == WD_ALIGN_PARAGRAPH.CENTER:
        return "center"
    if align == WD_ALIGN_PARAGRAPH.RIGHT:
        return "right"
    if align == WD_ALIGN_PARAGRAPH.JUSTIFY:
        return "justify"
    return None


# XML namespaces
_NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_NS_WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
_NS_A = "http://schemas.openxmlformats.org/drawingml/2006/main"
_NS_R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_NS_WP14 = "http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"


def _qname(tag: str, ns: str = _NS_W) -> str:
    return f"{{{ns}}}{tag}"


class DocxExtractor:
    def __init__(self, output_dir: str = ""):
        self.output_dir = output_dir
        self._img_counter: int = 0
        self._saved_images: Set[str] = set()

    def extract(self, docx_path: str) -> Document:
        doc = DocxDocument(docx_path)
        elements: List[BlockElement] = []
        list_buffer: List[ListItem] = []
        list_ordered: bool = False
        list_tight: bool = True
        code_buffer: List[str] = []

        footnotes = self._load_footnotes(doc)
        comments = self._load_comments(doc)

        for child in doc.element.body:
            tag = child.tag

            if tag.endswith("}p"):
                para = self._find_paragraph_in_doc(doc, child)
                if para is None:
                    continue

                # Page break
                if self._has_page_break(child):
                    self._flush(elements, code_buffer, list_buffer, list_ordered, list_tight)
                    elements.append(PageBreak())
                    list_buffer = []
                    list_tight = True
                    code_buffer = []
                    continue

                # Horizontal rule
                if self._is_horizontal_rule(para):
                    self._flush(elements, code_buffer, list_buffer, list_ordered, list_tight)
                    elements.append(HorizontalRule())
                    list_buffer = []
                    list_tight = True
                    code_buffer = []
                    continue

                # Extract runs (text, hyperlinks, images) from XML
                runs, para_images = self._extract_runs_and_images(child, doc)

                # Check monospace → code block
                flat_text = "".join(r.text for r in runs if isinstance(r, TextRun))
                is_code = runs and all(isinstance(r, TextRun) and r.code for r in runs) and flat_text.strip()
                if is_code:
                    code_buffer.append(flat_text)
                    continue
                elif code_buffer:
                    self._flush_code_buffer(elements, code_buffer)
                    code_buffer = []

                # Merge images into runs if no runs
                if not runs and para_images:
                    runs = para_images

                element = self._extract_paragraph_element(para, runs)

                if isinstance(element, ListBlock):
                    items = element.items
                    if items:
                        if list_buffer and list_ordered == element.ordered:
                            list_buffer.extend(items)
                        else:
                            self._flush_list_buffer(elements, list_buffer, list_ordered, list_tight)
                            list_buffer = list(items)
                            list_ordered = element.ordered
                            list_tight = element.tight
                    continue

                self._flush_list_buffer(elements, list_buffer, list_ordered, list_tight)
                list_buffer = []
                list_tight = True

                if element is not None:
                    elements.append(element)

                # Footnotes & comments
                elements.extend(self._get_para_footnotes(child, footnotes))
                elements.extend(self._get_para_comments(child, comments))

            elif tag.endswith("}tbl"):
                self._flush(elements, code_buffer, list_buffer, list_ordered, list_tight)
                list_buffer = []
                code_buffer = []
                table_elem = self._extract_table_xml(child, doc)
                if table_elem is not None:
                    elements.append(table_elem)

        self._flush(elements, code_buffer, list_buffer, list_ordered, list_tight)
        return Document(elements=elements)

    def _flush(self, elements, code_buffer, list_buffer, list_ordered, list_tight):
        self._flush_code_buffer(elements, code_buffer)
        self._flush_list_buffer(elements, list_buffer, list_ordered, list_tight)

    def _flush_code_buffer(self, elements, code_buffer):
        if code_buffer:
            elements.append(CodeBlock(code="\n".join(code_buffer)))

    def _flush_list_buffer(self, elements, buffer, ordered, tight):
        if buffer:
            elements.append(ListBlock(ordered=ordered, items=buffer, tight=tight))

    def _find_paragraph_in_doc(self, doc: DocxDocument, xml_elem):
        for para in doc.paragraphs:
            if para._p is xml_elem:
                return para
        return None

    def _has_page_break(self, p_elem) -> bool:
        try:
            tree = etree.fromstring(etree.tostring(p_elem))
            for br in tree.iter(_qname("br")):
                if br.get(_qname("type")) == "page":
                    return True
        except Exception:
            pass
        return False

    def _extract_paragraph_element(self, para, runs) -> Optional[BlockElement]:
        style_name = para.style.name if para.style else ""

        if style_name.startswith("Heading"):
            try:
                level = int(style_name.split()[-1])
            except (ValueError, IndexError):
                level = 1
            return Heading(level=level, runs=runs)

        if style_name in ("List Bullet", "List Number") or \
           style_name.startswith("List Bullet ") or \
           style_name.startswith("List Number "):
            ordered = "Number" in style_name
            item = ListItem(elements=[Paragraph(runs=runs)])
            tight = not para.paragraph_format.space_before \
                    and not para.paragraph_format.space_after
            return ListBlock(ordered=ordered, items=[item], tight=tight)

        numPr = para._p.find(qn("w:pPr"))
        if numPr is not None:
            numPr = numPr.find(qn("w:numPr"))
        if numPr is not None:
            item = ListItem(elements=[Paragraph(runs=runs)])
            return ListBlock(ordered=True, items=[item], tight=True)

        alignment = _alignment_to_str(para.alignment)
        return Paragraph(runs=runs, alignment=alignment)

    def _extract_runs_and_images(self, p_elem, doc) -> (List[InlineElement], List[Image]):
        runs: List[InlineElement] = []
        images: List[Image] = []

        try:
            tree = etree.fromstring(etree.tostring(p_elem))
        except Exception:
            return runs, images

        # Direct children only
        for child in tree:
            tag = child.tag

            if tag == _qname("r"):
                text = self._get_run_text(child)
                rPr = child.find(_qname("rPr"))

                if text:
                    runs.append(self._build_text_run(text, rPr))

                # Check for image in drawing inside this run
                drawing_elem = child.find(_qname("drawing"))
                if drawing_elem is not None:
                    img_elem = drawing_elem.find(f"{{{_NS_WP}}}inline")
                    if img_elem is None:
                        img_elem = drawing_elem.find(f"{{{_NS_WP14}}}anchor")
                    if img_elem is not None:
                        img = self._extract_single_image(img_elem, doc)
                        if img:
                            images.append(img)

            elif tag == _qname("hyperlink"):
                hyperlink = self._extract_hyperlink(child, doc)
                if hyperlink:
                    runs.append(hyperlink)

        return runs, images

    def _extract_hyperlink(self, hyperlink_elem, doc) -> Optional[Hyperlink]:
        r_id = hyperlink_elem.get(_qname("id", _NS_R))
        url = ""
        if r_id:
            try:
                rel = doc.part.rels[r_id]
                url = rel.target_ref if rel else ""
            except (KeyError, AttributeError):
                pass

        hyper_runs: List[TextRun] = []
        for r in hyperlink_elem.iter(_qname("r")):
            text = self._get_run_text(r)
            if not text:
                continue
            rPr = r.find(_qname("rPr"))
            hyper_runs.append(self._build_text_run(text, rPr))

        if not hyper_runs:
            return None
        return Hyperlink(url=url, runs=hyper_runs)

    def _get_run_text(self, run_elem) -> str:
        texts: List[str] = []
        for t in run_elem.iter(_qname("t")):
            if t.text:
                texts.append(t.text)
        return "".join(texts)

    def _build_text_run(self, text: str, rPr) -> TextRun:
        bold = False
        italic = False
        underline = False
        strike = False
        superscript = False
        subscript = False
        font_name = None
        font_size_pt = None

        if rPr is not None:
            b_elem = rPr.find(_qname("b"))
            if b_elem is not None:
                bold = b_elem.get(_qname("val")) != "false"

            i_elem = rPr.find(_qname("i"))
            if i_elem is not None:
                italic = i_elem.get(_qname("val")) != "false"

            u_elem = rPr.find(_qname("u"))
            if u_elem is not None:
                u_val = u_elem.get(_qname("val"))
                underline = u_val and u_val != "none"

            s_elem = rPr.find(_qname("strike"))
            if s_elem is not None:
                strike = s_elem.get(_qname("val")) != "false"

            va = rPr.find(_qname("vertAlign"))
            if va is not None:
                val = va.get(_qname("val"))
                if val == "superscript":
                    superscript = True
                elif val == "subscript":
                    subscript = True

            rFonts = rPr.find(_qname("rFonts"))
            if rFonts is not None:
                font_name = (
                    rFonts.get(_qname("ascii")) or
                    rFonts.get(_qname("hAnsi")) or
                    rFonts.get(_qname("eastAsia")) or
                    rFonts.get(_qname("cs"))
                )

            sz = rPr.find(_qname("sz"))
            if sz is not None:
                sz_val = sz.get(_qname("val"))
                if sz_val:
                    font_size_pt = int(sz_val) // 2

        code = _is_mono_font(font_name)
        return TextRun(
            text=text, bold=bold, italic=italic, code=code,
            underline=underline, strikethrough=strike,
            superscript=superscript, subscript=subscript,
            font_name=font_name, font_size=font_size_pt,
        )

    def _is_horizontal_rule(self, para) -> bool:
        try:
            pPr = para._p.find(qn("w:pPr"))
            if pPr is not None:
                pBdr = pPr.find(qn("w:pBdr"))
                if pBdr is not None:
                    bottom = pBdr.find(qn("w:bottom"))
                    if bottom is not None and bottom.get(qn("w:val")) == "single":
                        return True
        except Exception:
            pass
        return False

    # --- Images ---
    def _extract_single_image(self, drawing, doc) -> Optional[Image]:
        extent = drawing.find(_qname("extent", _NS_WP))
        cx = int(extent.get("cx")) if extent is not None and extent.get("cx") else None
        cy = int(extent.get("cy")) if extent is not None and extent.get("cy") else None

        docPr = drawing.find(_qname("docPr", _NS_WP))
        alt_text = ""
        if docPr is not None:
            alt_text = docPr.get("descr") or docPr.get("name") or ""

        blip = drawing.find(f".//{_qname('blip', _NS_A)}")
        if blip is None:
            return None
        embed = blip.get(_qname("embed", _NS_R))
        if not embed:
            return None

        try:
            image_part = doc.part.related_parts[embed]
        except KeyError:
            return None
        except AttributeError:
            return None

        img_name = self._save_image(image_part)
        rel_path = "images/" + img_name

        width_str = f"{cx / 914400:.2f}in" if cx else None
        height_str = f"{cy / 914400:.2f}in" if cy else None
        return Image(src=rel_path, alt=alt_text, width=width_str, height=height_str)

    def _save_image(self, image_part) -> str:
        ext = self._get_image_ext(image_part)
        self._img_counter += 1
        filename = f"image_{self._img_counter:03d}{ext}"
        self._saved_images.add(filename)
        if self.output_dir:
            filepath = os.path.join(self.output_dir, filename)
            with open(filepath, "wb") as f:
                f.write(image_part.blob)
        return filename

    @staticmethod
    def _get_image_ext(image_part) -> str:
        ct = image_part.content_type or ""
        mapping = {
            "image/png": ".png", "image/jpeg": ".jpg",
            "image/gif": ".gif", "image/bmp": ".bmp",
            "image/tiff": ".tiff", "image/svg+xml": ".svg",
        }
        for c, ext in mapping.items():
            if c in ct:
                return ext
        return ".png"

    # --- Tables ---
    def _extract_table_xml(self, tbl_elem, doc: DocxDocument) -> Optional[Table]:
        try:
            tree = etree.fromstring(etree.tostring(tbl_elem))
        except Exception:
            return None

        rows_xml = tree.findall(f".//{_qname('tr')}")
        if not rows_xml:
            return None

        all_rows: List[List[str]] = []
        aligns: List[Optional[str]] = []

        for row_idx, row_xml in enumerate(rows_xml):
            cells = row_xml.findall(f".//{_qname('tc')}")
            row_data: List[str] = []
            for cell_idx, cell in enumerate(cells):
                texts: List[str] = []
                for p in cell.findall(f".//{_qname('p')}"):
                    for t in p.findall(f".//{_qname('t')}"):
                        if t.text:
                            texts.append(t.text)
                    texts.append("\n")
                cell_text = "".join(texts).strip()
                row_data.append(cell_text)

                if row_idx == 0:
                    p = cell.find(f".//{_qname('p')}")
                    align_val = None
                    if p is not None:
                        pPr = p.find(_qname("pPr"))
                        if pPr is not None:
                            jc = pPr.find(_qname("jc"))
                            if jc is not None:
                                val = jc.get(_qname("val"))
                                align_map = {"left": "left", "center": "center",
                                             "right": "right", "both": "justify"}
                                align_val = align_map.get(val)
                    while len(aligns) <= cell_idx:
                        aligns.append(None)
                    aligns[cell_idx] = align_val

            all_rows.append(row_data)

        if not all_rows:
            return None

        num_cols = max(len(r) for r in all_rows)
        headers = all_rows[0] if all_rows else []
        body_rows = all_rows[1:] if len(all_rows) > 1 else []

        aligns = aligns[:num_cols] if aligns else []
        while len(aligns) < num_cols:
            aligns.append(None)

        return Table(headers=headers, rows=body_rows, align=aligns)

    # --- Footnotes & Comments ---
    def _load_footnotes(self, doc) -> dict:
        fn_map = {}
        try:
            part = doc.part.package.part_related_by(
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"
            )
            tree = etree.fromstring(part.blob)
            for fn in tree.findall(f".//{_qname('footnote')}"):
                fn_id = fn.get(_qname("id"))
                texts: List[str] = []
                for t in fn.iter(_qname("t")):
                    if t.text:
                        texts.append(t.text)
                fn_map[fn_id] = "".join(texts)
        except Exception:
            pass
        return fn_map

    def _get_para_footnotes(self, p_elem, fn_map: dict) -> List[Footnote]:
        results = []
        try:
            tree = etree.fromstring(etree.tostring(p_elem))
            for ref in tree.iter(_qname("footnoteReference")):
                fn_id = ref.get(_qname("id"))
                if fn_id and fn_id in fn_map:
                    results.append(Footnote(footnote_id=fn_id, text=fn_map[fn_id]))
        except Exception:
            pass
        return results

    def _load_comments(self, doc) -> dict:
        cm_map = {}
        try:
            part = doc.part.package.part_related_by(
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
            )
            tree = etree.fromstring(part.blob)
            for cm in tree.findall(f".//{_qname('comment')}"):
                cm_id = cm.get(_qname("id"), "")
                author = cm.get(_qname("author"), "")
                date_str = cm.get(_qname("date"), "")
                texts: List[str] = []
                for t in cm.iter(_qname("t")):
                    if t.text:
                        texts.append(t.text)
                cm_map[cm_id] = {
                    "author": author,
                    "text": "".join(texts),
                    "date": date_str or None,
                }
        except Exception:
            pass
        return cm_map

    def _get_para_comments(self, p_elem, cm_map: dict) -> List[Comment]:
        results = []
        try:
            tree = etree.fromstring(etree.tostring(p_elem))
            for ref in tree.iter(_qname("commentReference")):
                cm_id = ref.get(_qname("id"))
                if cm_id and cm_id in cm_map:
                    info = cm_map[cm_id]
                    results.append(Comment(author=info["author"], text=info["text"],
                                           date=info["date"]))
        except Exception:
            pass
        return results
